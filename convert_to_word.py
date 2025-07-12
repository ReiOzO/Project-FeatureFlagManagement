#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def markdown_to_word(md_file, docx_file):
    # Đọc file Markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Tạo document mới
    doc = Document()
    
    # Thiết lập styles
    styles = doc.styles
    
    # Style cho heading 1
    h1_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.space_after = Pt(12)
    
    # Style cho heading 2
    h2_style = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.size = Pt(16)
    h2_style.font.bold = True
    h2_style.space_after = Pt(10)
    
    # Style cho heading 3
    h3_style = styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
    h3_style.font.size = Pt(14)
    h3_style.font.bold = True
    h3_style.space_after = Pt(8)
    
    # Style cho normal text
    normal_style = styles.add_style('Normal Text', WD_STYLE_TYPE.PARAGRAPH)
    normal_style.font.size = Pt(12)
    normal_style.space_after = Pt(6)
    
    # Style cho code
    code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.size = Pt(10)
    code_style.font.name = 'Courier New'
    code_style.space_after = Pt(6)
    
    # Tách nội dung thành các dòng
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            doc.add_paragraph()
            continue
            
        # Heading 1
        if line.startswith('# '):
            text = line[2:]
            p = doc.add_paragraph(text, style='Heading 1')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Heading 2
        elif line.startswith('## '):
            text = line[3:]
            doc.add_paragraph(text, style='Heading 2')
            
        # Heading 3
        elif line.startswith('### '):
            text = line[4:]
            doc.add_paragraph(text, style='Heading 3')
            
        # Code block
        elif line.startswith('```'):
            continue
            
        # List item
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            p = doc.add_paragraph(text, style='Normal Text')
            p.paragraph_format.left_indent = Inches(0.25)
            
        # Numbered list
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            p = doc.add_paragraph(text, style='Normal Text')
            p.paragraph_format.left_indent = Inches(0.25)
            
        # Bold text
        elif '**' in line:
            # Xử lý bold text
            parts = line.split('**')
            p = doc.add_paragraph(style='Normal Text')
            for i, part in enumerate(parts):
                if i % 2 == 0:  # Regular text
                    if part:
                        p.add_run(part)
                else:  # Bold text
                    run = p.add_run(part)
                    run.bold = True
                    
        # Regular paragraph
        else:
            doc.add_paragraph(line, style='Normal Text')
    
    # Lưu document
    doc.save(docx_file)
    print(f"Đã tạo file Word: {docx_file}")

if __name__ == "__main__":
    markdown_to_word('Project_Proposal_Feature_Flag_Management.md', 'Project_Proposal_Feature_Flag_Management.docx') 