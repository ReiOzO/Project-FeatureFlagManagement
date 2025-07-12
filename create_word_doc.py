#!/usr/bin/env python3
# -*- coding: utf-8 -*-

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    print("Thư viện python-docx đã được cài đặt")
except ImportError:
    print("Thư viện python-docx chưa được cài đặt. Vui lòng chạy: pip install python-docx")
    exit(1)

def create_project_proposal():
    # Tạo document mới
    doc = Document()
    
    # Thiết lập styles
    styles = doc.styles
    
    # Style cho heading 1
    h1_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.space_after = Pt(12)
    
    # Style cho heading 2
    h2_style = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.size = Pt(16)
    h2_style.font.bold = True
    h2_style.space_after = Pt(10)
    
    # Style cho heading 3
    h3_style = styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
    h3_style.font.size = Pt(14)
    h3_style.font.bold = True
    h3_style.space_after = Pt(8)
    
    # Style cho normal text
    normal_style = styles.add_style('Normal Text', WD_STYLE_TYPE.PARAGRAPH)
    normal_style.font.size = Pt(12)
    normal_style.space_after = Pt(6)
    
    # Tiêu đề chính
    title = doc.add_paragraph("Đề xuất Dự án: Hệ thống Quản lý Feature Flag với Custom Solution và AWS AppConfig", style='Heading 1')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Tổng quan Dự án
    doc.add_paragraph("1. Tổng quan Dự án", style='Heading 2')
    
    # 1.1. Giới thiệu
    doc.add_paragraph("1.1. Giới thiệu", style='Heading 3')
    doc.add_paragraph("Dự án này đề xuất việc phát triển một hệ thống Feature Flag Management hoàn chỉnh, kết hợp giữa custom solution và AWS AppConfig để quản lý tính năng trong ứng dụng web. Hệ thống hỗ trợ gradual rollouts, A/B testing, automated rollback, và performance monitoring.", style='Normal Text')
    
    doc.add_paragraph("Feature Flag (còn gọi là Feature Toggle) là một kỹ thuật phát triển phần mềm cho phép các nhà phát triển bật/tắt các tính năng từ xa mà không cần triển khai lại mã nguồn. Điều này giúp giảm thiểu rủi ro khi phát hành các tính năng mới và cho phép kiểm soát chính xác hơn đối với trải nghiệm người dùng.", style='Normal Text')
    
    # 1.2. Vấn đề cần giải quyết
    doc.add_paragraph("1.2. Vấn đề cần giải quyết", style='Heading 3')
    doc.add_paragraph("Trong quá trình phát triển phần mềm hiện đại, các tổ chức gặp phải nhiều thách thức:", style='Normal Text')
    
    problems = [
        "Chu kỳ phát hành chậm: Việc triển khai các tính năng mới thường đòi hỏi toàn bộ ứng dụng phải được xây dựng lại và triển khai lại",
        "Rủi ro phát hành cao: Khi phát hành tính năng mới, nếu có lỗi, toàn bộ ứng dụng có thể bị ảnh hưởng",
        "Thiếu khả năng kiểm soát: Khó khăn trong việc phát hành tính năng cho một nhóm người dùng cụ thể hoặc theo tỷ lệ phần trăm",
        "Thiếu dữ liệu về hiệu suất: Khó đánh giá tác động của tính năng mới đối với hiệu suất hệ thống và trải nghiệm người dùng"
    ]
    
    for problem in problems:
        p = doc.add_paragraph(problem, style='Normal Text')
        p.paragraph_format.left_indent = Inches(0.25)
    
    # 1.3. Mục tiêu dự án
    doc.add_paragraph("1.3. Mục tiêu dự án", style='Heading 3')
    doc.add_paragraph("Dự án này nhằm mục đích xây dựng một hệ thống Feature Flag Management toàn diện để giải quyết các vấn đề trên thông qua:", style='Normal Text')
    
    objectives = [
        "Triển khai cơ bản Feature Flag: Xây dựng hệ thống cho phép bật/tắt tính năng từ xa",
        "Gradual Rollouts: Hỗ trợ triển khai dần dần theo phần trăm người dùng",
        "A/B Testing: Cung cấp khả năng kiểm thử A/B với các biến thể khác nhau",
        "Automated Rollback: Tự động phát hiện vấn đề và rollback khi cần thiết",
        "Performance Monitoring: Giám sát hiệu suất real-time của các tính năng",
        "User Segmentation: Phân chia người dùng thành các nhóm để nhắm mục tiêu cụ thể",
        "Analytics: Thu thập và phân tích dữ liệu về việc sử dụng tính năng"
    ]
    
    for i, objective in enumerate(objectives, 1):
        p = doc.add_paragraph(f"{i}. {objective}", style='Normal Text')
        p.paragraph_format.left_indent = Inches(0.25)
    
    # 2. Phạm vi Dự án
    doc.add_paragraph("2. Phạm vi Dự án", style='Heading 2')
    
    # 2.1. Các thành phần chính
    doc.add_paragraph("2.1. Các thành phần chính", style='Heading 3')
    doc.add_paragraph("Dự án sẽ bao gồm các thành phần chính sau:", style='Normal Text')
    
    components = [
        "Backend Service: API RESTful, tích hợp AWS AppConfig, hệ thống caching, WebSocket",
        "Frontend Dashboard: Giao diện quản lý feature flags, bảng điều khiển metrics, công cụ A/B testing",
        "AWS Integration: AWS AppConfig, CloudWatch, Lambda functions",
        "Documentation: API documentation, hướng dẫn triển khai, tài liệu người dùng"
    ]
    
    for i, component in enumerate(components, 1):
        p = doc.add_paragraph(f"{i}. {component}", style='Normal Text')
        p.paragraph_format.left_indent = Inches(0.25)
    
    # 2.2. Những gì không thuộc phạm vi
    doc.add_paragraph("2.2. Những gì không thuộc phạm vi", style='Heading 3')
    doc.add_paragraph("Để đảm bảo dự án hoàn thành trong thời gian 4 tuần, những nội dung sau sẽ không nằm trong phạm vi:", style='Normal Text')
    
    exclusions = [
        "Tích hợp với các hệ thống CI/CD bên ngoài",
        "Hỗ trợ đa ngôn ngữ trong giao diện người dùng",
        "Mobile application cho quản lý feature flags",
        "Tích hợp với các hệ thống phân tích dữ liệu bên ngoài",
        "Sử dụng AWS S3 cho lưu trữ dữ liệu"
    ]
    
    for exclusion in exclusions:
        p = doc.add_paragraph(f"• {exclusion}", style='Normal Text')
        p.paragraph_format.left_indent = Inches(0.25)
    
    # 3. Phân tích Kỹ thuật
    doc.add_paragraph("3. Phân tích Kỹ thuật", style='Heading 2')
    
    # 3.1. Kiến trúc hệ thống
    doc.add_paragraph("3.1. Kiến trúc hệ thống", style='Heading 3')
    doc.add_paragraph("Hệ thống sẽ được xây dựng với kiến trúc 3 tầng:", style='Normal Text')
    
    architecture = [
        "Frontend (React): Giao diện người dùng, biểu đồ, WebSocket client",
        "Backend (Node.js): API RESTful, logic xử lý, tích hợp AWS",
        "AWS Services: AppConfig, CloudWatch, Lambda functions"
    ]
    
    for arch in architecture:
        p = doc.add_paragraph(f"• {arch}", style='Normal Text')
        p.paragraph_format.left_indent = Inches(0.25)
    
    # 3.2. Công nghệ sử dụng
    doc.add_paragraph("3.2. Công nghệ sử dụng", style='Heading 3')
    
    tech_backend = [
        "Node.js và Express.js",
        "AWS SDK cho tích hợp với AWS services",
        "WebSocket cho real-time updates",
        "In-memory caching"
    ]
    
    doc.add_paragraph("Backend:", style='Normal Text')
    for tech in tech_backend:
        p = doc.add_paragraph(f"• {tech}", style='Normal Text')
        p.paragraph_format.left_indent = Inches(0.5)
    
    tech_frontend = [
        "React.js",
        "Tailwind CSS cho UI",
        "Chart.js cho data visualization",
        "WebSocket client"
    ]
    
    doc.add_paragraph("Frontend:", style='Normal Text')
    for tech in tech_frontend:
        p = doc.add_paragraph(f"• {tech}", style='Normal Text')
        p.paragraph_format.left_indent = Inches(0.5)
    
    tech_aws = [
        "AWS AppConfig: Quản lý cấu hình feature flags",
        "CloudWatch: Monitoring và metrics",
        "Lambda: Automated rollback functions"
    ]
    
    doc.add_paragraph("AWS Services:", style='Normal Text')
    for tech in tech_aws:
        p = doc.add_paragraph(f"• {tech}", style='Normal Text')
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 3.3. Cấu trúc dữ liệu Feature Flag
    doc.add_paragraph("3.3. Cấu trúc dữ liệu Feature Flag", style='Heading 3')
    doc.add_paragraph("Hệ thống sẽ sử dụng cấu trúc dữ liệu JSON để lưu trữ thông tin feature flags, bao gồm các thành phần chính sau:", style='Normal Text')
    
    data_structure = [
        "Thông tin phiên bản: Version number và timestamp cập nhật cuối cùng",
        "Cấu hình Feature Flag: Trạng thái kích hoạt, phần trăm rollout, user targeting, variants, metadata",
        "Cấu trúc variants: Hỗ trợ multiple variants cho A/B testing với weight distribution",
        "Targeting rules: Nhắm mục tiêu người dùng dựa trên user groups, specific user IDs"
    ]
    
    for structure in data_structure:
        p = doc.add_paragraph(f"• {structure}", style='Normal Text')
        p.paragraph_format.left_indent = Inches(0.25)
    
    # 3.4. API Endpoints
    doc.add_paragraph("3.4. API Endpoints", style='Heading 3')
    doc.add_paragraph("Hệ thống sẽ cung cấp RESTful API với các nhóm endpoints chính sau:", style='Normal Text')
    
    api_groups = [
        "Feature Flag Management: Quản lý danh sách, cấu hình rollout, đánh giá feature flag, batch operations",
        "Analytics và Metrics: Metrics tổng quan, metrics chi tiết, A/B testing results",
        "System Health và Monitoring: Health check, AWS integration status, performance metrics",
        "Cache Management: Cache refresh, cache status"
    ]
    
    for i, api_group in enumerate(api_groups, 1):
        p = doc.add_paragraph(f"{i}. {api_group}", style='Normal Text')
        p.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_paragraph("Tất cả API endpoints sẽ tuân theo chuẩn RESTful, sử dụng HTTP methods phù hợp và trả về responses theo format JSON thống nhất.", style='Normal Text')
    
    # 4. Kế hoạch Thực hiện
    doc.add_paragraph("4. Kế hoạch Thực hiện", style='Heading 2')
    
    # 4.1. Lịch trình dự án (4 tuần)
    doc.add_paragraph("4.1. Lịch trình dự án (4 tuần)", style='Heading 3')
    
    weeks = [
        "Tuần 1: Thiết lập và Cấu hình Cơ bản - Thiết lập môi trường, cấu hình AWS AppConfig, xây dựng backend cơ bản",
        "Tuần 2: Phát triển Core Features - Hoàn thiện API endpoints, xây dựng caching, tích hợp AWS AppConfig",
        "Tuần 3: Frontend và Monitoring - Phát triển giao diện, tích hợp CloudWatch, xây dựng cảnh báo",
        "Tuần 4: Testing và Hoàn thiện - Kiểm thử toàn diện, tối ưu hóa, hoàn thiện tài liệu"
    ]
    
    for i, week in enumerate(weeks, 1):
        p = doc.add_paragraph(f"Tuần {i}: {week}", style='Normal Text')
        p.paragraph_format.left_indent = Inches(0.25)
    
    # 4.2. Phân công công việc
    doc.add_paragraph("4.2. Phân công công việc", style='Heading 3')
    doc.add_paragraph("Dự án sẽ được thực hiện bởi một thành viên duy nhất với các vai trò đa dạng:", style='Normal Text')
    
    roles = [
        "Quản lý dự án tổng thể",
        "Phát triển backend services",
        "Phát triển frontend dashboard",
        "Tích hợp AWS services",
        "Kiểm thử hệ thống",
        "Viết tài liệu"
    ]
    
    doc.add_paragraph("Full-stack Developer:", style='Normal Text')
    for role in roles:
        p = doc.add_paragraph(f"• {role}", style='Normal Text')
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 5. Phân tích Lợi ích
    doc.add_paragraph("5. Phân tích Lợi ích", style='Heading 2')
    
    # 5.1. Lợi ích kỹ thuật
    doc.add_paragraph("5.1. Lợi ích kỹ thuật", style='Heading 3')
    
    tech_benefits = [
        "Tăng tốc độ phát triển: Tách biệt việc triển khai mã và kích hoạt tính năng",
        "Giảm thiểu rủi ro: Khả năng rollback nhanh chóng khi phát hiện vấn đề",
        "Tăng cường kiểm soát: Kiểm soát chính xác đối tượng người dùng nhìn thấy tính năng",
        "Cải thiện quy trình phát triển: Hỗ trợ trunk-based development"
    ]
    
    for benefit in tech_benefits:
        p = doc.add_paragraph(f"• {benefit}", style='Normal Text')
        p.paragraph_format.left_indent = Inches(0.25)
    
    # 5.2. Lợi ích kinh doanh
    doc.add_paragraph("5.2. Lợi ích kinh doanh", style='Heading 3')
    
    business_benefits = [
        "Tăng tốc độ ra thị trường: Phát hành tính năng nhanh hơn",
        "Cải thiện trải nghiệm người dùng: A/B testing để tối ưu hóa trải nghiệm",
        "Tối ưu hóa quyết định: Thu thập dữ liệu về hiệu suất tính năng",
        "Giảm chi phí: Giảm thời gian phát triển và chi phí khắc phục lỗi"
    ]
    
    for benefit in business_benefits:
        p = doc.add_paragraph(f"• {benefit}", style='Normal Text')
        p.paragraph_format.left_indent = Inches(0.25)
    
    # 6. Quản lý Rủi ro
    doc.add_paragraph("6. Quản lý Rủi ro", style='Heading 2')
    
    # 6.1. Rủi ro tiềm ẩn
    doc.add_paragraph("6.1. Rủi ro tiềm ẩn", style='Heading 3')
    
    risks = [
        "Thời gian phát triển không đủ (4 tuần là khá ngắn cho 1 người)",
        "Thiếu kinh nghiệm với AWS AppConfig",
        "Quá tải công việc do chỉ có 1 người thực hiện",
        "Vấn đề về hiệu suất khi số lượng feature flags tăng lên"
    ]
    
    for risk in risks:
        p = doc.add_paragraph(f"• {risk}", style='Normal Text')
        p.paragraph_format.left_indent = Inches(0.25)
    
    # 6.2. Chiến lược giảm thiểu rủi ro
    doc.add_paragraph("6.2. Chiến lược giảm thiểu rủi ro", style='Heading 3')
    
    strategies = [
        "Ưu tiên các tính năng cốt lõi",
        "Tập trung vào MVP (Minimum Viable Product) trước",
        "Tận dụng tài liệu và best practices của AWS",
        "Áp dụng phương pháp phát triển linh hoạt"
    ]
    
    for strategy in strategies:
        p = doc.add_paragraph(f"• {strategy}", style='Normal Text')
        p.paragraph_format.left_indent = Inches(0.25)
    
    # 7. Yêu cầu Tài nguyên
    doc.add_paragraph("7. Yêu cầu Tài nguyên", style='Heading 2')
    
    # 7.1. Nhân sự
    doc.add_paragraph("7.1. Nhân sự", style='Heading 3')
    doc.add_paragraph("• 1 Full-stack Developer (thực hiện tất cả các vai trò)", style='Normal Text')
    
    # 7.2. Công nghệ và Công cụ
    doc.add_paragraph("7.2. Công nghệ và Công cụ", style='Heading 3')
    
    tools = [
        "Phát triển: Visual Studio Code, Git và GitHub, Postman",
        "AWS Services: AWS AppConfig, AWS CloudWatch, AWS Lambda",
        "Môi trường: Local development environment, AWS Free Tier account"
    ]
    
    for tool in tools:
        p = doc.add_paragraph(f"• {tool}", style='Normal Text')
        p.paragraph_format.left_indent = Inches(0.25)
    
    # 7.3. Chi phí dự kiến
    doc.add_paragraph("7.3. Chi phí dự kiến", style='Heading 3')
    doc.add_paragraph("Dự án sẽ được thực hiện trong phạm vi AWS Free Tier nên không phát sinh chi phí.", style='Normal Text')
    
    # 8. Kết luận và Đề xuất
    doc.add_paragraph("8. Kết luận và Đề xuất", style='Heading 2')
    
    # 8.1. Tóm tắt
    doc.add_paragraph("8.1. Tóm tắt", style='Heading 3')
    doc.add_paragraph("Dự án Feature Flag Management với Custom Solution và AWS AppConfig đề xuất xây dựng một hệ thống toàn diện để quản lý việc phát hành tính năng, hỗ trợ gradual rollouts, A/B testing, và automated rollback.", style='Normal Text')
    
    # 8.2. Đề xuất
    doc.add_paragraph("8.2. Đề xuất", style='Heading 3')
    doc.add_paragraph("Chúng tôi đề xuất phê duyệt dự án này với thời gian thực hiện 4 tuần và không phát sinh chi phí (sử dụng AWS Free Tier). Dự án sẽ được thực hiện bởi một thành viên duy nhất với vai trò full-stack developer, tập trung vào việc xây dựng MVP trước và mở rộng tính năng sau.", style='Normal Text')
    
    # 8.3. Các bước tiếp theo
    doc.add_paragraph("8.3. Các bước tiếp theo", style='Heading 3')
    
    next_steps = [
        "Phê duyệt đề xuất dự án",
        "Thiết lập môi trường phát triển",
        "Cấu hình AWS services",
        "Bắt đầu phát triển theo lịch trình đã đề xuất"
    ]
    
    for i, step in enumerate(next_steps, 1):
        p = doc.add_paragraph(f"{i}. {step}", style='Normal Text')
        p.paragraph_format.left_indent = Inches(0.25)
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph("Dự án này được đề xuất như một phần của chương trình thực tập tại AWS, nhằm mục đích chứng minh khả năng triển khai giải pháp cloud phức tạp sử dụng AWS services.", style='Normal Text')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Lưu document
    doc.save('Project_Proposal_Feature_Flag_Management.docx')
    print("Đã tạo file Word: Project_Proposal_Feature_Flag_Management.docx")

if __name__ == "__main__":
    create_project_proposal() 